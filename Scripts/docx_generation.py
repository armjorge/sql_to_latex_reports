import os
from docx import Document
import pandas as pd
from sqlalchemy import create_engine
from colorama import Fore, Style, init
from datetime import datetime
from dotenv import load_dotenv
import yaml

class DOCX_TO_LATEX:
    def __init__(self, working_folder, data_access, queries_folder, template_file):
        self.working_folder = working_folder
        self.data_access = data_access 
        self.queries_folder = queries_folder
        self.template_file = template_file
        self.output_folder = os.path.join(self.working_folder, 'Reportes BI')

    def sql_conexion(self):
        sql_url = self.data_access['sql_url']
        #url example: 'postgresql://arXXXrge:XXX@ep-shy-darkness-10211313-poolXXXX.tech/neondb?sslmode=require&channel_binding=require'
        try:
            engine = create_engine(sql_url)
            return engine
        except Exception as e:
            print(f"❌ Error connecting to database: {e}")
            return None
        
    def sql_to_dataframe(self): 
        engine = self.sql_conexion()  # must return a SQLAlchemy Engine
        if engine is None:
            print("❌ No se pudo obtener el engine de SQL.")
            return False
        dataframes = {}
        for file in os.listdir(self.queries_folder):
            if file.endswith(".sql"):
                name = os.path.splitext(file)[0]
                print(f"{name}")
                with open(os.path.join(self.queries_folder, file)) as f:
                    query = f.read()
                df = pd.read_sql(query, engine)
                dataframes[name] = df
        print(f"{Fore.GREEN} Extracción de dataframes finalizada. {Style.RESET_ALL}")
        return dataframes

    def generate_report(self, dataframes):
        # Abrir template
        doc = Document(self.template_file)

        # Formato de fecha en español
        meses = {
            1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
            5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
            9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
        }
        hoy = datetime.now()
        fecha_texto = f"{hoy.day} de {meses[hoy.month]} de {hoy.year}"

        # Recorrer cada párrafo y reemplazar placeholders
        for paragraph in doc.paragraphs:
            text = paragraph.text

            # Reemplazo de la fecha
            if "{date}" in text:
                text = text.replace("{date}", fecha_texto)

            # Reemplazo de los dataframes
            for key, df in dataframes.items():
                placeholder = f"{{{key}.sql}}"
                if placeholder in text:
                    text_block = self._df_to_text(df)
                    text = text.replace(placeholder, text_block)

            paragraph.text = text        
        print(f"\t🧩 Generando documento DOCX basado en template: {self.template_file}")



        # Recorrer cada párrafo y reemplazar el placeholder
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for key, df in dataframes.items():
                placeholder = f"{{{key}.sql}}"
                if placeholder in text:
                    # Convertir el DataFrame a texto limpio
                    text_block = self._df_to_text(df)
                    new_text = text.replace(placeholder, text_block)
                    paragraph.text = new_text

        # Crear carpeta de salida
        os.makedirs(self.output_folder, exist_ok=True)
        output_filename = datetime.now().strftime('%Y-%m-%d %Hh%Mm Seguimiento.docx')
        output_path = os.path.join(self.output_folder, output_filename)

        # Guardar documento final
        doc.save(output_path)
        print(f"\t✅ Reporte guardado en: {output_path}")
        return output_path

    def _df_to_text(self, df: pd.DataFrame) -> str:
        """
        Convierte un DataFrame a texto plano legible, con formato monetario en columnas 'importe'
        y sin incluir encabezados.
        """
        if df.empty:
            return "(Sin registros)"

        df_copy = df.copy()

        # Formatear columna 'file_date' si es la única columna
        if list(df_copy.columns) == ['file_date']:
            try:
                ts = pd.to_datetime(df_copy['file_date'].iloc[0])
                return ts.strftime("%d/%m/%Y %Hh%Mm")
            except Exception as e:
                return f"(Error formateando fecha: {e})"

        # Detectar columnas relacionadas con 'importe' y formatearlas como moneda
        for col in df_copy.columns:
            if 'importe' in col.lower():
                df_copy[col] = pd.to_numeric(df_copy[col], errors='coerce').apply(
                    lambda x: f"${x:,.2f}" if pd.notna(x) else ""
                )

        # Convertir filas a texto (sin encabezados)
        text_lines = []
        for _, row in df_copy.iterrows():
            row_text =  "\t" + " | ".join(str(x) for x in row.values)
            row_text = row_text.replace("| None |", " | ")
            text_lines.append(row_text)

        return "\n".join(text_lines)

    def reporting_docx_run(self):
        print(f"{Fore.CYAN}\t⚡ Ejecutando consultas SQL para generación de reportes DOCX...{Style.RESET_ALL}")
        dataframes = self.sql_to_dataframe()
        for key, value in dataframes.items(): 
            print(key)
            print(value.head())
        self.generate_report(dataframes)
        print(f"{Fore.GREEN}\t✅ Reporte DOCX generado exitosamente.{Style.RESET_ALL}")


if __name__ == "__main__":
    # Set working folder
    script_dir = os.path.dirname(os.path.abspath(__file__))
    root = os.path.abspath(os.path.join(script_dir, ".."))
    env_file = os.path.join(root, ".env")
    # 1. Load .env if it exists
    if os.path.exists(env_file):
        load_dotenv(env_file)
    folder_root = os.getenv("MAIN_PATH")
    working_folder = os.path.join(folder_root, "Local Data")
    # Set the dict with sql URL
    yaml_path = os.path.join(working_folder, "config.yaml")
    yaml_exists = os.path.exists(yaml_path)
    integration_path = os.path.join(working_folder, "Integración")
    if yaml_exists:
        # Abrir y cargar el contenido YAML en un diccionario
        with open(yaml_path, 'r', encoding='utf-8') as f:
            data_access = yaml.safe_load(f)
        print(f"✅ Archivo YAML cargado correctamente: {os.path.basename(yaml_path)}")   

    # Set queries folder
    queries_folder = os.path.join(root, "sql_queries")
    # Initialize and run
    template_file = os.path.join('.', 'Templates', 'template.docx') # Root, templates. 
    app = DOCX_TO_LATEX(working_folder, data_access, queries_folder, template_file)
    app.reporting_docx_run()    